import docx

doc_path = r"C:\Flutter\guvnl_project\Document\SEM5_report.docx"
doc = docx.Document(doc_path)

print("Document paragraphs:")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        # Only print headings or the first few paragraphs
        if para.style.name.startswith('Heading') or i < 20:
            print(f"Para {i} [{para.style.name}]: {para.text.strip()[:100]}")

print("\nDocument tables:")
for i, table in enumerate(doc.tables):
    print(f"Table {i}: {len(table.rows)} rows, {len(table.columns)} columns")
