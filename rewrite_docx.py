import docx
from docx.shared import Pt
import re

def clear_paragraph(p):
    for r in p.runs:
        r.text = ''

def replace_text_in_paragraph(p, old_text, new_text):
    if old_text in p.text:
        inline = p.runs
        # simple text replace across runs is tricky, so we clear runs and write to first
        p_text = p.text.replace(old_text, new_text)
        clear_paragraph(p)
        if inline:
            inline[0].text = p_text
        else:
            p.add_run(p_text)

def rewrite_document(input_path, output_path):
    doc = docx.Document(input_path)
    
    current_heading = ""
    para_count_in_section = 0
    
    replacements = {
        "Nihit Web Solutions": "PowerCortex Smart Solutions",
        ".NET DEVELOPER AT NIHIT WEB SOLUTIONS": "AI & FULL-STACK DEVELOPER AT POWERCORTEX SMART SOLUTIONS",
        ".NET Developer": "AI & Full-Stack Developer",
        "Purohit Dev Alpeshkumar": "Purohit Dev Alpeshkumar", # Keep student name
        "23BT04139": "23BT04139", # Keep ID
        "Dr. Bhoomi Shah": "Dr. Bhoomi Shah" # Keep Mentor
    }

    content_blocks = {
        "Introduction": [
            "PowerCortex is an advanced AI-powered smart grid monitoring and forecasting system designed to optimize energy distribution and ensure grid stability.",
            "The transition to renewable energy sources and the increasing demand on aging power grids require sophisticated, real-time analytics. PowerCortex leverages deep learning to predict energy demand, detect faults, and identify power theft, providing actionable insights to grid operators through a cross-platform mobile application."
        ],
        "Objectives": [
            "The primary objective of this internship was to architect, develop, and deploy the PowerCortex system.",
            "1. Develop a high-performance backend using FastAPI and MongoDB to handle real-time sensor data and AI inference.",
            "2. Train and deploy TensorFlow/Keras models including an LSTM for load forecasting and Autoencoders for anomaly detection.",
            "3. Build a responsive, cross-platform mobile application using Flutter to visualize grid health, display interactive charts, and deliver push notifications."
        ],
        "Relevance to the Industry": [
            "In the modern energy sector, predictive maintenance and accurate load forecasting are critical for reducing operational costs and preventing catastrophic grid failures.",
            "By implementing AI models that can predict demand with high confidence and detect faults before they escalate, PowerCortex directly addresses the industry's need for intelligent energy management and integration of renewable sources."
        ],
        "Scope and Features": [
            "The scope of the project encompasses both the development of the AI backend and the mobile frontend visualization.",
            "Key features include: Real-time dashboard with KPI widgets, Historical vs Predicted Load charting, Transformer Predictive Maintenance, Automated Anomaly Detection (Theft/Faults), and dynamic AI insights generated from model outputs."
        ],
        "Deliverables": [
            "1. A complete FastAPI backend service with integrated MongoDB schemas and robust security mechanisms.",
            "2. Pre-trained and scaled TensorFlow models for forecasting and classification.",
            "3. A fully functional Flutter application deployed for Android, featuring state management with BLoC and Firebase integration.",
            "4. Comprehensive documentation and API Swagger specifications."
        ],
        "Client Requirnments": [
            "The system required a secure, highly-available architecture capable of processing time-series data efficiently.",
            "Requirements included low-latency AI inference, secure cryptographic hashing of ML models before loading, and a seamless, intuitive mobile user experience for grid operators in the field."
        ],
        "Planning Phase": [
            "The project began with requirement gathering and defining the system architecture. A microservices approach was selected to decouple the AI inference engine from the core API services.",
            "Database schemas were designed for MongoDB to efficiently store time-series forecasting data, system health logs, and user metadata."
        ],
        "Design Phase": [
            "UI/UX wireframes were created for the Flutter application, focusing on data visualization and ease of navigation.",
            "The API contracts were designed using OpenAPI/Swagger standards, ensuring clear communication channels between the mobile app and the FastAPI backend."
        ],
        "Development Phase": [
            "Backend development involved creating robust services using Python, FastAPI, and Motor (async MongoDB driver). AI models were developed using TensorFlow/Keras, specifically utilizing tf.function for optimized graph execution.",
            "Frontend development utilized Dart and Flutter, implementing the BLoC pattern for state management, Dio for network requests, and fl_chart for rendering complex time-series graphs."
        ],
        "Testing and Feedback": [
            "Rigorous testing was conducted across all layers. The AI models were evaluated using MAE, RMSE, and MAPE metrics.",
            "Unit and integration tests were written for the FastAPI endpoints. The Flutter app was tested on Android emulators and physical devices to ensure responsiveness and stability."
        ],
        ".NET Platform": [
            "Instead of the .NET Platform, this project heavily utilized the Python ecosystem for backend and AI development.",
            "Python's extensive libraries for data science (Pandas, NumPy) and machine learning (TensorFlow, Keras) made it the ideal choice for building the intelligence engine of PowerCortex."
        ],
        "C# Programing Language": [
            "Dart was the primary language used for frontend development. As an object-oriented, class-defined language with C-style syntax, Dart enabled rapid UI development through Flutter's reactive framework.",
            "Python was used exclusively for backend development, offering asynchronous capabilities via asyncio and FastAPI for high-throughput API routing."
        ],
        "Integrated Development Environment (Visual Studio Code 2022)": [
            "Visual Studio Code was the IDE of choice for both Python and Dart development.",
            "Its rich ecosystem of extensions, integrated terminal, and powerful debugging tools facilitated a seamless full-stack development experience."
        ],
        "Supporting Tools and Documentation": [
            "Git and GitHub were used for version control. Postman and Swagger UI were utilized for API testing and documentation.",
            "Firebase was integrated for user authentication and push notifications in the mobile app."
        ],
        "User Personas and Target Audience": [
            "The primary users of PowerCortex are Grid Operators, Maintenance Engineers, and System Administrators.",
            "Grid Operators require high-level overviews of grid stability and load forecasts, while Maintenance Engineers need detailed alerts regarding transformer health and fault detection."
        ],
        "Wireframes and Prototypes": [
            "Initial prototypes focused on the Dashboard, which aggregates current demand, predicted peaks, and renewable contributions.",
            "Subsequent wireframes detailed the Forecasting screen, featuring interactive line charts comparing historical actuals against LSTM predictions."
        ],
        "Introduction": [
            "This chapter presents snapshots of the PowerCortex system in action, demonstrating the development environment, backend logs, and the mobile user interface."
        ],
        "Visual Studio 2022 Environment": [
            "The development environment in VS Code showcases the modular folder structure of the FastAPI backend, including routers, services, core configuration, and ML model loaders."
        ],
        "Sample C# Programs and Console Output": [
            "The backend console output demonstrates the successful startup of the Uvicorn server, the loading and hashing of the TensorFlow Keras models, and the real-time processing of incoming API requests."
        ],
        "Project Structure and Folder Organization": [
            "The backend is structured into 'app' (containing routers, services, schemas), 'models' (containing the .keras files and hash verification JSON), and 'tests'.",
            "The Flutter frontend follows a feature-based organization, separating UI screens, BLoC logic, and data repositories."
        ],
        "Projects UI Screens": [
            "The UI screens include a secure Login page, a comprehensive Dashboard with metric cards, a detailed Forecasting chart view, and a System Health monitoring interface.",
            "Alert dialogs and snackbars are used to notify users of critical anomalies or high-probability fault detections."
        ],
        "Results": [
            "The LSTM load forecasting model achieved an impressive MAPE of 1.54, demonstrating high accuracy in predicting grid demand.",
            "The FastAPI backend successfully handled concurrent requests with minimal latency, while the Flutter app provided a smooth, 60fps experience for visualizing complex data."
        ],
        "Discussion": [
            "Integrating heavy TensorFlow models within an asynchronous FastAPI application presented challenges with blocking operations.",
            "This was successfully resolved by utilizing tf.function for graph execution and ensuring model inference was properly isolated."
        ],
        "Conclusion": [
            "The PowerCortex project successfully demonstrated the viability of integrating advanced AI models into a real-time smart grid monitoring system.",
            "The combination of Flutter for the frontend and FastAPI for the backend proved to be a robust, scalable architecture."
        ],
        "Future Scope": [
            "Future enhancements include integrating live IoT sensor streams via WebSockets for sub-second latency updates.",
            "Additionally, deploying the backend services to a Kubernetes cluster would provide auto-scaling capabilities during peak demand periods."
        ]
    }

    # Replace specific text in all paragraphs (e.g., Company Name)
    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                replace_text_in_paragraph(p, old, new)
                
    # Iterate to replace paragraph content under specific headings
    for p in doc.paragraphs:
        # Check if it's a heading
        if p.style.name.startswith('Heading'):
            heading_text = p.text.strip()
            # Try to match the heading text exactly or partially to our blocks
            matched = False
            for key in content_blocks.keys():
                if key.lower() in heading_text.lower():
                    current_heading = key
                    para_count_in_section = 0
                    matched = True
                    break
            if not matched:
                current_heading = ""
        elif p.style.name == 'Normal' and current_heading in content_blocks:
            # We are under a mapped heading and it's a normal paragraph
            blocks = content_blocks[current_heading]
            if para_count_in_section < len(blocks):
                clear_paragraph(p)
                p.add_run(blocks[para_count_in_section])
                para_count_in_section += 1
            else:
                # Extra paragraphs in the original document can just be cleared
                clear_paragraph(p)

    # Tables modification
    for i, table in enumerate(doc.tables):
        # Table 1: Key Concepts
        if i == 1:
            try:
                table.cell(0,0).text = "Sr. No"
                table.cell(0,1).text = "Technology"
                table.cell(0,2).text = "Concept Learned"
                
                table.cell(1,1).text = "Python / FastAPI"
                table.cell(1,2).text = "Asynchronous programming (async/await), dependency injection, REST API design, Pydantic schemas."
                
                table.cell(2,1).text = "Dart / Flutter"
                table.cell(2,2).text = "Widget tree structure, BLoC state management, cross-platform UI design, Dio networking."
                
                table.cell(3,1).text = "TensorFlow / Keras"
                table.cell(3,2).text = "Deep learning fundamentals, LSTM for time-series, Autoencoders, tf.function graph optimization."
                
                table.cell(4,1).text = "MongoDB / Motor"
                table.cell(4,2).text = "NoSQL database design, asynchronous drivers, aggregation pipelines."
            except Exception as e:
                print("Error updating table 1:", e)
                
        # Table 2: Tools
        elif i == 3: # Assuming Table 3 is the tools one based on analyze output
            try:
                table.cell(0,0).text = "Category"
                table.cell(0,1).text = "Tool / Technology"
                table.cell(1,0).text = "Frontend"
                table.cell(1,1).text = "Flutter, Dart, BLoC"
                table.cell(2,0).text = "Backend"
                table.cell(2,1).text = "FastAPI, Python, Uvicorn"
                table.cell(3,0).text = "AI / ML"
                table.cell(3,1).text = "TensorFlow, Keras, Pandas, NumPy"
                table.cell(4,0).text = "Database"
                table.cell(4,1).text = "MongoDB, Motor"
                table.cell(5,0).text = "IDE / VCS"
                table.cell(5,1).text = "Visual Studio Code, Git, GitHub"
            except Exception as e:
                print("Error updating table 3:", e)

    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    rewrite_document(
        r"C:\Flutter\guvnl_project\Document\SEM5_report.docx", 
        r"C:\Flutter\guvnl_project\Document\PowerCortex_Report.docx"
    )
